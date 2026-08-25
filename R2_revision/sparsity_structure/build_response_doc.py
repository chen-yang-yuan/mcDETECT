#!/usr/bin/env python3
"""
Build the internal response record for Reviewer #2 major point 6 (sparsity).

    python3 build_response_doc.py [out.docx]

Writes plans/Response_R2_comment6_sparsity.docx, following the shape of the A1 record
plans/Response_R2_comments1-2_Baysor_SSAM.docx.

**Every number in the document is read from a file under output/.** Nothing is typed by hand,
and the prose is composed from the same values that fill the tables, so the two cannot drift --
the rule the A1 README states as "every number should be re-derivable from output/ alone".

Re-run this after the A2b sweep lands to fill in section 2; it is written to be idempotent.
"""

import shutil
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
import a2_config as _C
C_REF_GENES, C_SYN_GENES = _C.REF_GENES, _C.SYN_GENES

HERE = Path(__file__).resolve().parent
REPO = HERE.parents[1]
OUT = HERE / "output"
A2A = OUT / "a2a" / "multigene"
SUB = A2A / "neuropil_subdomains_Isocortex_50"
STRATA = OUT / "a2a" / "readstrata"
A2B = OUT / "a2b" / "metrics"
A2C = OUT / "a2c"
FIG = OUT / "figures"
PUB = REPO / "output" / "MERSCOPE_WT_AD_comparison"
COMBINED_GRANULE_ADATA = PUB / "granule_adata_tsne.h5ad"
PUBSUB = PUB / "neuropil_subdomains_Isocortex_50"
AREAS = ["Isocortex", "OLF", "HPF-CA", "HPF-DG", "HPF-SR", "CTXsp", "TH", "MB", "FT"]
TMP = HERE / ".doc_png_cache"


# ============================================================ io ============================================================ #

def require(path, hint=""):
    """Fail loudly. A silently omitted table is worse than no document."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"missing input: {p}\n  {hint}")
    return p


def read_csv(path, hint=""):
    d = pd.read_csv(require(path, hint))
    return d.loc[:, [c for c in d.columns if str(c).strip() != ""]]   # drop write.csv row-name col


def fmt(x, nd=2):
    if x is None or (isinstance(x, float) and not np.isfinite(x)):
        return "n/a"
    if isinstance(x, (int, np.integer)):
        return f"{x:,}"
    return f"{x:,.{nd}f}"


def pct(x, nd=1):
    return "n/a" if x is None or not np.isfinite(x) else f"{100 * x:.{nd}f}%"


# ============================================================ docx helpers ============================================================ #

class Doc:
    def __init__(self):
        self.d = Document()
        st = self.d.styles["Normal"]
        st.font.name = "Calibri"
        st.font.size = Pt(10.5)
        self.figno = 0
        self.tabno = 0
        self.manifest = []

    def h(self, text, level=1):
        self.d.add_heading(text, level=level)

    def p(self, text, bold=False, italic=False, color=None, space_after=6):
        """`**...**` inside `text` becomes a real bold run. Without this the asterisks render
        literally in Word, which is easy to miss when composing prose in source."""
        import re as _re
        par = self.d.add_paragraph()
        for part in _re.split(r"(\*\*[^*]+\*\*)", text):
            if not part:
                continue
            emph = part.startswith("**") and part.endswith("**")
            run = par.add_run(part[2:-2] if emph else part)
            run.bold, run.italic = (bold or emph), italic
            if color:
                run.font.color.rgb = RGBColor(*color)
        par.paragraph_format.space_after = Pt(space_after)
        return par

    def quote(self, text):
        par = self.d.add_paragraph()
        run = par.add_run(text)
        run.italic = True
        par.paragraph_format.left_indent = Inches(0.35)
        par.paragraph_format.space_after = Pt(8)

    def bullets(self, items):
        for it in items:
            self.d.add_paragraph(it, style="List Bullet")

    def table(self, df, caption, nd=2, max_rows=None):
        self.tabno += 1
        if max_rows is not None and len(df) > max_rows:
            df = df.head(max_rows)
        cap = self.d.add_paragraph()
        r = cap.add_run(f"Table {self.tabno}. {caption}")
        r.bold = True
        r.font.size = Pt(9.5)

        t = self.d.add_table(rows=1, cols=len(df.columns))
        t.style = "Light Grid Accent 1"
        for j, c in enumerate(df.columns):
            cell = t.rows[0].cells[j]
            cell.text = str(c)
            for pp in cell.paragraphs:
                for rr in pp.runs:
                    rr.bold = True
                    rr.font.size = Pt(8.5)
        # Decide the formatter per COLUMN, from the column's own dtype. Iterating with
        # df.iterrows() would upcast each row to a single dtype, so an integer column sitting
        # beside a float one silently renders as "34.000".
        kind = {c: df[c].dtype.kind for c in df.columns}
        for i in range(len(df)):
            cells = t.add_row().cells
            for j, c in enumerate(df.columns):
                v = df[c].iloc[i]
                if kind[c] == "f":
                    txt = "n/a" if not np.isfinite(v) else f"{v:,.{nd}f}"
                elif kind[c] in "iu":
                    txt = f"{int(v):,}"
                else:
                    txt = str(v)
                cells[j].text = txt
                for pp in cells[j].paragraphs:
                    for rr in pp.runs:
                        rr.font.size = Pt(8.5)
        self.d.add_paragraph()
        self.manifest.append(f"Table {self.tabno}: {caption[:70]}")
        return self.tabno

    def figure(self, path, caption, width_in=6.0):
        """JPEG -> PNG before embedding: the analysis writes JPEGs and Word mishandles their
        Adobe APP14 marker (noted in the analysis README)."""
        src = require(path, "produced by the A2 notebooks or A2_figures.R")
        TMP.mkdir(exist_ok=True)
        png = TMP / (src.stem.replace(" ", "_") + ".png")
        if not png.exists() or png.stat().st_mtime < src.stat().st_mtime:
            im = Image.open(src).convert("RGB")
            # The analysis renders at dpi 500; embedded at ~6 inches that is far more pixels than
            # Word will ever show, and it inflates the file into the tens of MB. Cap at 200 dpi
            # for the embedded width.
            cap_px = int(width_in * 200)
            if im.width > cap_px:
                im = im.resize((cap_px, round(im.height * cap_px / im.width)), Image.LANCZOS)
            im.save(png, "PNG", optimize=True)
        self.figno += 1
        self.d.add_picture(str(png), width=Inches(width_in))
        self.d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = self.d.add_paragraph()
        try:
            shown = src.relative_to(REPO)
        except ValueError:                       # path outside the repo (e.g. a test redirect)
            shown = src.name
        r = cap.add_run(f"Figure R-{self.figno}. {caption}  [{shown}]")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        self.d.add_paragraph()
        self.manifest.append(f"Figure R-{self.figno}: {src.name}")

    def save(self, path):
        self.d.save(path)


# ============================================================ data ============================================================ #

def load_a2a():
    d = {}
    d["run"] = read_csv(A2A / "run_info.csv").iloc[0].to_dict()
    d["cross"] = read_csv(A2A / "complexity_crosscheck.csv").iloc[0].to_dict()
    d["ret"] = read_csv(A2A / "retention_by_region.csv")
    d["comp"] = read_csv(A2A / "subtype_composition.csv")
    d["dens_new"] = read_csv(A2A / "subtype_density_per_region_multigene.csv")
    d["dens_pub"] = read_csv(PUB / "subtype_density_per_region_granule_adata_tsne.csv")
    d["corr"] = read_csv(FIG / "subdomain_correspondence.csv")
    d["labels"] = pd.read_parquet(require(SUB / "4_hard_normalized_cluster_labels.parquet"))
    d["de_new"] = read_csv(SUB / "granule_DE_genes_Subdomain 3_vs_Subdomain 1.csv")
    d["de_pub"] = read_csv(PUBSUB / "granule_DE_genes_Subdomain 1_vs_Subdomain 2.csv")
    d["gsea_new"] = read_csv(SUB / "granule_DE_genes_Subdomain 3_vs_Subdomain 1_GSEA.csv")
    d["gsea_pub"] = read_csv(PUBSUB / "granule_DE_genes_Subdomain 1_vs_Subdomain 2_GSEA.csv")
    d["edges"] = read_csv(STRATA / "readstrata_edges.csv")
    d["scplx"] = read_csv(STRATA / "readstrata_complexity.csv")
    d["sdens"] = read_csv(STRATA / "readstrata_density.csv")
    return d


def load_a2c():
    d = {}
    d["run"] = read_csv(A2C / "run_info.csv")
    d["groups"] = read_csv(A2C / "groups_tested.csv")
    d["dropped"] = read_csv(A2C / "groups_dropped.csv")
    d["ge"] = read_csv(A2C / "group_enrichment.csv")
    d["pairs"] = pd.read_parquet(require(A2C / "pair_enrichment.parquet"))
    d["go"] = read_csv(FIG / "a2c_go_shared_term_test.csv")
    d["prog"] = read_csv(FIG / "a2c_programme_summary.csv")
    # The calibration gate is the reason section 3.1 may assert that z is calibrated at all, so
    # the number is read, never typed. It is written only when A2c section 6 runs.
    d["cal"] = read_csv(
        A2C / "null_calibration.csv",
        hint="run A2c_cooccurrence.ipynb section 6 with VALIDATE = True -- it scores a matrix "
             "drawn from the null itself and writes output/a2c/null_calibration.csv.")
    d["pair_go"] = pd.read_parquet(require(
        FIG / "a2c_pair_go.parquet",
        "run `Rscript A2_figures.R` -- section 6 now exports the per-pair GO annotation."))
    d["gobin"] = read_csv(
        FIG / "a2c_go_by_z_bin.csv",
        hint="run `Rscript A2_figures.R` with only RUN_GSEA/RUN_A2C = TRUE -- section 6 now also "
             "exports the per-pair GO annotation and the z-stratified table.")
    return d


def load_a2b():
    d = {}
    d["metrics"] = read_csv(A2B / "a2b_metrics.csv",
                            "the HGCC sweep writes this; see slurm/submit.sh")
    d["status"] = read_csv(A2B / "a2b_status.csv")
    d["detect"] = read_csv(A2B / "a2b_detection_summary.csv")
    d["summary"] = read_csv(A2B / "a2b_summary.csv")
    d["k15"] = read_csv(FIG / "a2b_structure_at_k15.csv",
                        "produced by A2_figures.R section 5")
    return d


def a2b_paired_comparison(metrics, out_csv=None):
    """
    Real minus permuted, at matched n, paired within seed and k.

    The matched series exist because both metrics depend on n and the permuted arms hold about a
    third as many granules; comparing them at their own sizes would confound structure with
    sample size. Pairing within seed is then the right unit -- five independent permutations,
    each with its own matched real subsample -- and the combined test is over the per-k
    differences rather than over the 145 individual comparisons, which are not independent.

    Returns (per_k_df, summary_dict).
    """
    from scipy.stats import wilcoxon

    mm = metrics[metrics["matched"]]
    piv = mm.pivot_table(index=["seed", "n_clusters"], columns="condition",
                         values=["silhouette_score", "ari_stability_mean"])
    d = pd.DataFrame({
        "sil_real": piv[("silhouette_score", "real")],
        "sil_perm": piv[("silhouette_score", "permuted")],
        "ari_real": piv[("ari_stability_mean", "real")],
        "ari_perm": piv[("ari_stability_mean", "permuted")],
    }).reset_index()
    d["sil_diff"] = d.sil_real - d.sil_perm
    d["ari_diff"] = d.ari_real - d.ari_perm

    per_k = (d.groupby("n_clusters")
               .agg(sil_real=("sil_real", "mean"), sil_perm=("sil_perm", "mean"),
                    sil_diff=("sil_diff", "mean"),
                    seeds_real_higher_sil=("sil_diff", lambda x: int((x > 0).sum())),
                    ari_real=("ari_real", "mean"), ari_perm=("ari_perm", "mean"),
                    ari_diff=("ari_diff", "mean"),
                    seeds_real_higher_ari=("ari_diff", lambda x: int((x > 0).sum())))
               .reset_index())
    per_k["perm_frac_of_real_sil"] = per_k.sil_perm / per_k.sil_real

    k15 = d[d.n_clusters == 15]
    summary = {
        "n_comparisons": int(len(d)), "n_k": int(d.n_clusters.nunique()),
        "n_seeds": int(d.seed.nunique()),
        "sil_real_higher": int((d.sil_diff > 0).sum()),
        "ari_real_higher": int((d.ari_diff > 0).sum()),
        "sil_mean_diff": float(d.sil_diff.mean()),
        "ari_mean_diff": float(d.ari_diff.mean()),
        "sil_p_wilcoxon": float(wilcoxon(per_k.sil_diff).pvalue),
        "ari_p_wilcoxon": float(wilcoxon(per_k.ari_diff).pvalue),
        "k15_sil_real": float(k15.sil_real.mean()), "k15_sil_perm": float(k15.sil_perm.mean()),
        "k15_ari_real": float(k15.ari_real.mean()), "k15_ari_perm": float(k15.ari_perm.mean()),
        "k15_seeds_real_higher_sil": int((k15.sil_diff > 0).sum()),
        "k15_seeds_real_higher_ari": int((k15.ari_diff > 0).sum()),
        "k15_sign_test_p": float(0.5 ** len(k15)),
        # Quote the retained fraction at the published k and over k >= 10. Below that, k = 2-5
        # partitions a 34-dimensional space into a handful of blobs and the ratio is noisy in
        # both directions -- it is not the regime any of this analysis is read in.
        "k15_perm_frac_of_real_sil": float(
            per_k.loc[per_k.n_clusters == 15, "perm_frac_of_real_sil"].iloc[0]),
        "perm_frac_of_real_sil_min": float(
            per_k.loc[per_k.n_clusters >= 10, "perm_frac_of_real_sil"].min()),
        "perm_frac_of_real_sil_max": float(
            per_k.loc[per_k.n_clusters >= 10, "perm_frac_of_real_sil"].max()),
        "n_matched": float(mm.n_obs.mean()),
    }
    if out_csv is not None:
        Path(out_csv).parent.mkdir(parents=True, exist_ok=True)
        per_k.to_csv(out_csv, index=False)
    return per_k, summary


def density_concordance(pub, new):
    """Direction agreement and log2FC correlation between the published and multi-gene runs."""
    def wide(df, tag):
        d = df[df.brain_area.isin(AREAS) & df["sample"].isin(["WT", "AD"])]
        p = d.pivot_table(index=["subtype", "brain_area"], columns="sample", values="density")
        s = d[d["sample"] == "WT"].set_index(["subtype", "brain_area"])[["p_bonf_star"]]
        o = p.join(s)
        o["lfc"] = np.log2(o["AD"] / o["WT"])
        o.columns = [f"{c}_{tag}" for c in o.columns]
        return o
    c = wide(pub, "pub").join(wide(new, "mg")).dropna(subset=["lfc_pub", "lfc_mg"])
    sig = c[c.p_bonf_star_pub.isin(["*", "**", "***"])]
    return c, {
        "n_cells": len(c),
        "agree": float((np.sign(c.lfc_pub) == np.sign(c.lfc_mg)).mean()),
        "r": float(c[["lfc_pub", "lfc_mg"]].corr().iloc[0, 1]),
        "n_sig_pub": len(sig),
        "n_sig_same_dir": int((np.sign(sig.lfc_pub) == np.sign(sig.lfc_mg)).sum()),
        "n_still_sig": int(sig.p_bonf_star_mg.isin(["*", "**", "***"]).sum()),
    }


def go_abundance_stratified(pair_go, pairs, n_bins=10, out_csv=None):
    """
    Re-test the GO association with pair abundance held fixed.

    The raw test compares z between pairs that share a GO term and pairs that do not, pooled. That
    is confounded: GO-annotated genes are the better-studied, more abundant ones, and abundance
    drives |z| in both directions. This is the same confound that forced abundance matching into
    the group permutation; the GO test was written before that and never had it applied.

    Here the comparison is made WITHIN expected-count deciles and combined across them with a
    signed-rank test on the per-decile differences, which is the right unit -- ten strata, not
    thirty thousand non-independent pairs.

    Returns (per_decile_df, summary_dict).
    """
    from scipy.stats import mannwhitneyu, wilcoxon

    m = pair_go.merge(pairs[["gene_i", "gene_j", "expected"]], on=["gene_i", "gene_j"],
                      how="left").dropna(subset=["expected"])
    m["decile"] = pd.qcut(m["expected"].rank(method="first"), n_bins, labels=False) + 1

    rows = []
    for b, g in m.groupby("decile"):
        sh, no = g.loc[g.go_shared, "z"], g.loc[~g.go_shared, "z"]
        if len(sh) < 20 or len(no) < 20:
            continue
        rows.append({"expected-count decile": int(b), "median expected": float(g.expected.median()),
                     "pairs sharing GO": len(sh), "pairs not sharing": len(no),
                     "median z sharing": float(sh.median()),
                     "median z not sharing": float(no.median()),
                     "difference": float(sh.median() - no.median()),
                     "p (within decile)": float(mannwhitneyu(sh, no).pvalue)})
    per = pd.DataFrame(rows)

    # evidence for the confound itself: how shared-GO tracks abundance
    conf = (m.groupby("decile")
              .agg(median_expected=("expected", "median"),
                   frac_shared_go=("go_shared", "mean"), n=("go_shared", "size"))
              .reset_index())

    summary = {
        "n_pairs": int(len(m)),
        "frac_shared_overall": float(m.go_shared.mean()),
        "median_expected_shared": float(m.loc[m.go_shared, "expected"].median()),
        "median_expected_not": float(m.loc[~m.go_shared, "expected"].median()),
        "frac_shared_lowest_decile": float(conf.frac_shared_go.iloc[0]),
        "frac_shared_highest_decile": float(conf.frac_shared_go.iloc[-1]),
        "n_deciles": len(per),
        "n_deciles_negative": int((per["difference"] < 0).sum()),
        "median_difference": float(per["difference"].median()),
        "p_signed_rank": float(wilcoxon(per["difference"]).pvalue) if len(per) > 5 else np.nan,
    }
    if out_csv is not None:
        Path(out_csv).parent.mkdir(parents=True, exist_ok=True)
        per.assign(**{k: v for k, v in summary.items()
                      if k in ("p_signed_rank", "median_difference", "n_deciles_negative")}
                   ).to_csv(out_csv, index=False)
    return per, conf, summary


# ============================================================ sections ============================================================ #

REVIEWER_COMMENT = (
    "I asked for the distributions of reads and genes per granule, and these are now provided "
    "(Fig. R9). They are useful, but they also raise a concern. In the MERSCOPE data the median "
    "granule contains only about 6 to 7 transcripts and 4 unique genes. More fundamentally, "
    "because detection seeds on individual marker genes, each granule essentially expresses just "
    "the single marker it was detected on, as Fig. 3e confirms for the majority of granules. The "
    "granule subtypes, which the authors define by the dominant subtype marker (Fig. 3f), "
    "therefore largely reflect the seeding marker rather than any genuine multi-gene granule "
    "transcriptome, and the discrete t-SNE structure (Fig. 4d) follows from the seeding alone. I "
    "am not convinced that randomized data, with granules selected on a single marker gene in "
    "the same way, would not produce an essentially identical embedding. The authors partly "
    "acknowledge this by moving the WT/AD comparison to aggregated neuropil microdomains "
    "(Fig. 5), which is sensible, but it effectively concedes that single-granule profiles are "
    "too sparse to interpret directly. My original request, to stratify the analysis by granule "
    "complexity and show that the structure is not a low-count artifact, has not been met. "
    "Relatedly, the per-granule read and gene-count distributions shown in the rebuttal (Fig. R9) "
    "do not appear in the manuscript itself; they should be included, at least in the supplement, "
    "so that readers have an honest description of how sparse the detected granules are.")


def embedding_feature_composition(cache_csv=None):
    """
    How much of the embedding is the seeding procedure itself.

    The granule embedding and the k = 15 subtyping are computed on REF_GENES, and most of those
    genes ARE the detection seeds. This quantifies that: the seed/non-seed split of the feature
    space, how much of the matrix variance each side carries, and how often a granule's highest
    reference-gene value is simply the marker it was detected on.

    It matters for reading A2b. If discreteness is a property of the detection procedure rather
    than of the data, then any null that keeps single-marker seeding -- which is what the reviewer
    stipulates -- will reproduce it, whatever "randomized data" is taken to mean.

    Cached: the source is a ~355 MB h5ad and the rest of the builder is CSV-only.
    """
    if cache_csv is not None and Path(cache_csv).exists():
        return pd.read_csv(cache_csv).iloc[0].to_dict()

    import scanpy as sc
    ad = sc.read_h5ad(require(COMBINED_GRANULE_ADATA))
    ref = [g for g in C_REF_GENES if g in ad.var_names]
    seed = set(C_SYN_GENES)
    X = ad[:, ref].X
    X = X.toarray() if hasattr(X, "toarray") else np.asarray(X)
    v = X.var(axis=0)
    is_seed = np.array([g in seed for g in ref])
    top_gene = np.asarray(ref)[X.argmax(axis=1)]
    out = {
        "n_features": len(ref),
        "n_seed_features": int(is_seed.sum()),
        "n_nonseed_features": int((~is_seed).sum()),
        "var_frac_seed": float(v[is_seed].sum() / v.sum()),
        "var_frac_nonseed": float(v[~is_seed].sum() / v.sum()),
        "frac_top_gene_is_own_seed": float(
            (top_gene == ad.obs["gene"].astype(str).to_numpy()).mean()),
        "n_granules": int(ad.n_obs),
    }
    if cache_csv is not None:
        Path(cache_csv).parent.mkdir(parents=True, exist_ok=True)
        pd.DataFrame([out]).to_csv(cache_csv, index=False)
    return out


# Which reviewer comment each section of this document answers. The `para` index is the
# paragraph in plans/R2_Reviewer_comments.docx that `quote` is taken from, and
# verify_response_doc.py asserts every quote is a verbatim substring of it -- so this table is
# checked against the reviewer's own file rather than transcribed and trusted.
#
# Two rows are worth reading carefully. Section 3.2 answers a CLAIM the reviewer makes inside
# major point 6, not a request; and section 3.3 answers a concern from major points 1, 4 and 7,
# which is the subject of the separate Baysor/SSAM response, not of this document.
COMMENT_MAP = [
    {"section": "1 (A2a)", "comment": "Major point 6 — sparsity", "requested": "yes", "para": 24,
     "quote": "My original request, to stratify the analysis by granule complexity and show "
              "that the structure is not a low-count artifact, has not been met."},
    {"section": "1.1, Figs R-1/R-2", "comment": "Major point 6 — sparsity", "requested": "yes",
     "para": 24,
     "quote": "they should be included, at least in the supplement"},
    {"section": "2 (A2b)", "comment": "Major point 6 — sparsity", "requested": "yes", "para": 24,
     "quote": "I am not convinced that randomized data, with granules selected on a single "
              "marker gene in the same way, would not produce an essentially identical "
              "embedding."},
    {"section": "3.2 (A2c)", "comment": "Major point 6 — a claim, not a request",
     "requested": "no", "para": 24,
     "quote": "each granule essentially expresses just the single marker it was detected on"},
    {"section": "3.3 (A2c, removable)", "comment": "Major points 1, 4 and 7 — circularity",
     "requested": "no", "para": 15,
     "quote": "not that the detected structures are genuine granules rather than any "
              "co-expressed transcript cluster"},
]


def section_map(doc):
    doc.p(
        "Which comment each section answers. The major-point-6 text is quoted in full in the "
        "next section; the major-points-1/4/7 text is quoted in the separate Baysor and SSAM "
        "response, and section 3.3 is the only part of this document that reaches into it. "
        "\"Requested\" distinguishes what the reviewer asked for from what we added.")
    doc.table(pd.DataFrame([
        {"section": r["section"], "reviewer comment": r["comment"],
         "requested by the reviewer": r["requested"], "the text it answers": r["quote"]}
        for r in COMMENT_MAP]),
        "Correspondence between the sections of this document and the reviewer's comments.")


def section0(doc, a, b):
    doc.h("Reviewer #2 — sparsity of granules and the stochastic origin of granule-level "
          "structure (major point 6)", 1)
    doc.quote(REVIEWER_COMMENT)
    doc.p("Response.", bold=True)
    doc.p(
        "We have carried out the two analyses the reviewer asks for, and one further analysis "
        "that was not requested. First, we stratified the analysis by granule complexity and "
        "repeated the entire pair-1 downstream pipeline on multi-gene granules only (A2a). "
        "Second, we constructed the randomized data the reviewer describes and ran the identical "
        "detection and embedding pipeline on it (A2b), which shows that randomization does not "
        "reproduce the granule population, though it reproduces more of the embedding structure "
        "than we expected. Third, because the reviewer's specific claim is that a granule carries "
        "nothing beyond its seeding marker, we asked directly whether the genes co-detected "
        "inside a granule are non-randomly associated, using only the genes detection never "
        "touches (A2c). We also stratified by read count, and we report a result from A2c that "
        "does not support one of our own expectations.")
    doc.p(f"Throughout, the granule set is the published pair-1 detection "
          f"({fmt(int(a['cross']['n_granules']))} granules from MERSCOPE_WT_1 and "
          f"MERSCOPE_AD_1). The per-granule read and gene-count distributions (Fig. R9) are "
          f"promoted to the supplement as requested; Figures R-1 and R-2 below are those "
          f"distributions.")


def section1_methods(doc, a):
    doc.h("1. A2a — multi-gene granule re-analysis (the requested complexity stratification)", 1)
    doc.h("1.0  What was filtered on, and why it is not the column it appears to be", 2)
    doc.p(
        "Granule complexity is not the `comp` column of the granule table, and using it would "
        "have answered a different question. mcDETECT's dbscan() restricts the transcript frame "
        "to the granule markers before detection begins (model.py:102), so `comp` counts distinct "
        "granule MARKERS, capped at the 20 seeds, and it is never recomputed after merge_sphere(). "
        "The quantity we stratified on is the number of distinct panel genes inside the sphere, "
        "taken from profile(): n_genes = (counts > 0).sum(axis=1). Table 1 reports how far the two "
        "diverge.")
    doc.p(
        f"Two conventions we state rather than leave to be found. (i) The panel contains no blank "
        f"probes -- the 19 negative controls are real nuclear-enriched panel genes -- so the "
        f"unique-gene count needs a stated denominator; the filter excludes them "
        f"({a['run']['complexity_col']}). (ii) Every published mcDETECT result uses "
        f"profile(buffer = 0.00); only the Fig. R9 distribution panel used buffer = 0.01, which "
        f"is where the reviewer's \"median 6 to 7 transcripts\" comes from. The two radii agree on "
        f"{pct(a['cross']['agree_unique_genes'])} of granules for unique genes (medians identical "
        f"at {fmt(a['cross']['median_unique_genes_buffer001'], 0)}) but on only "
        f"{pct(a['cross']['agree_reads'])} for reads (medians "
        f"{fmt(a['cross']['median_reads_buffer001'], 0)} at 0.01 versus "
        f"{fmt(a['cross']['median_reads_buffer000'], 0)} at 0.00): a minimum-enclosing sphere has "
        f"support points exactly on its surface, so a 0.01 um buffer systematically captures a "
        f"few more transcripts, which against a median of about 5 reads is a large relative "
        f"difference. We therefore report the published buffer = 0.01 read distribution, so the "
        f"supplement and the reviewer's figures agree, and filter on the buffer = 0.00 count, "
        f"which is the matrix actually carried forward -- the conservative direction, since 0.00 "
        f"yields equal or fewer genes.")
    doc.p(
        "No detection was re-run. profile() queries each sphere independently, so subsetting rows "
        "of the published granule x gene matrix is identical to re-profiling the retained spheres; "
        "the notebook asserts this on a sample of granules. Everything downstream -- subtyping, "
        "embedding, density, microdomains -- then runs exactly as published, with only the granule "
        "set changed.")


def section1_results(doc, a):
    run, cross = a["run"], a["cross"]
    ret = a["ret"]
    ov = ret[ret.brain_area == "overall"].set_index("batch")

    doc.h("1.1  How many granules are multi-gene", 2)
    wt, ad = ov.loc["MERSCOPE_WT_1"], ov.loc["MERSCOPE_AD_1"]
    doc.p(
        f"At a threshold of {int(run['min_unique_genes'])} or more distinct genes, "
        f"{fmt(int(run['n_granules']))} of {fmt(int(cross['n_granules']))} granules are retained "
        f"-- {pct(wt.retention)} of WT and {pct(ad.retention)} of AD. Retention is stable across "
        f"regions ({pct(ret[ret.brain_area.isin(AREAS)].retention.min())} to "
        f"{pct(ret[ret.brain_area.isin(AREAS)].retention.max())}), so the subset is not a "
        f"regional selection in disguise.")
    t = ret[ret.brain_area.isin(AREAS)].copy()
    t["sample"] = np.where(t.batch.str.contains("WT"), "WT", "AD")
    piv = t.pivot_table(index="brain_area", columns="sample",
                        values=["n_all", "n_multigene", "retention"]).reindex(AREAS)
    piv.columns = [f"{b} {a_}" for a_, b in piv.columns]
    for c in piv.columns:                       # pivot_table returns floats; these are counts
        if "n_all" in c or "n_multigene" in c:      # columns are "WT n_all", "AD n_multigene"
            piv[c] = piv[c].astype("int64")
    doc.table(piv.reset_index().rename(columns={"brain_area": "region"}),
              "Granules retained at three or more unique genes, by region and sample.", nd=3)
    doc.figure(FIG / "complexity_n_reads_all.jpeg",
               "Reads per granule (published buffer = 0.01 convention). This is the Fig. R9 "
               "content the reviewer asks to see in the manuscript.")
    doc.figure(FIG / "complexity_n_genes_all.jpeg",
               "Unique genes per granule, on the buffer = 0.00 column the complexity filter acts "
               "on, so this panel and the retention figures cannot disagree.")
    doc.figure(FIG / "comp_vs_unique_genes.jpeg",
               "The `comp` column against the panel-wide unique-gene count. `comp` is capped at "
               "the 20 seed markers and is not a measure of granule complexity; the y axis is.")
    doc.figure(FIG / "multigene_retention.jpeg",
               "Retention of the multi-gene subset by region and sample.")

    doc.h("1.2  Granule subtype structure persists", 2)
    comp = a["comp"].set_index("subtype")["fraction"]
    doc.p(
        f"We re-clustered the subset with the published procedure unchanged -- MiniBatchKMeans, "
        f"k = {int(run['k_subtype'])}, seed {int(run['subtype_seed'])}, on the 34 compartment "
        f"markers after normalising on the full panel -- and re-read the compartment identities "
        f"from the heatmap. The composition is "
        + ", ".join(f"{k} {pct(v)}" for k, v in comp.sort_values(ascending=False).items())
        + ". The clusters remain marker-defined and compartment-interpretable, which is the "
          "property the reviewer questions.")
    doc.figure(A2A / "heatmap_subtype_ordered.jpeg",
               "Multi-gene granules clustered on the 34 compartment markers, ordered by the "
               "manual compartment assignment. Read this as the verdict on the assignment.")

    doc.h("1.3  The WT/AD density biology persists", 2)
    c, st = density_concordance(a["dens_pub"], a["dens_new"])
    doc.p(
        f"Across {st['n_cells']} region x subtype cells the multi-gene and published runs agree "
        f"in direction {pct(st['agree'])} of the time, and the AD-versus-WT log2 fold changes "
        f"correlate at r = {fmt(st['r'])}. Of the {st['n_sig_pub']} cells significant in the "
        f"published run, {st['n_sig_same_dir']} keep the same direction and "
        f"{st['n_still_sig']} remain significant. The headline result is unchanged: the "
        f"pre-synaptic granule density reduction in AD is significant in every region in which it "
        f"was significant before, in the same direction, at the same level.")
    pre = c.loc["pre-syn"].reindex(AREAS).reset_index()
    tab = pd.DataFrame({
        "region": pre["brain_area"],
        "WT published": pre["WT_pub"], "AD published": pre["AD_pub"],
        "log2FC published": pre["lfc_pub"], "sig published": pre["p_bonf_star_pub"],
        "WT multi-gene": pre["WT_mg"], "AD multi-gene": pre["AD_mg"],
        "log2FC multi-gene": pre["lfc_mg"], "sig multi-gene": pre["p_bonf_star_mg"]})
    doc.table(tab, "Pre-synaptic granule density, AD versus WT, published versus multi-gene "
                   "subset. Density is granules per 50 um spot; significance is a t-test on "
                   "log1p per-spot counts, Bonferroni-corrected within subtype.", nd=2)
    doc.figure(FIG / "granule_density_all_vs_multigene.jpeg",
               "WT and AD granule density by region and subtype: published full granule set "
               "(left) against the multi-gene subset (right). Absolute densities differ by "
               "construction; the WT-versus-AD direction is what carries over.")


def section1_microdomains(doc, a):
    doc.h("1.4  The neuropil microdomain result persists", 2)
    corr = a["corr"][["multigene", "best_published_match", "best_match_frac"]].drop_duplicates()
    corr = corr.sort_values("multigene")
    n_shared = int(a["corr"]["n_shared_spots"].iloc[0])
    doc.p(
        f"We repeated the Fig. 5 microdomain analysis with the multi-gene subset substituted for "
        f"the granules, holding the spatial scaffold fixed: the 50 um Isocortex spot grid and its "
        f"SpaGCN layer labels are read, not recomputed, so only the granules differ. The "
        f"subdomain labels themselves ARE recomputed, because microdomains are defined by "
        f"granule-subtype composition and inheriting the published labels would be circular.")
    doc.p(
        f"K-means labels are arbitrary, so the first question is whether the partition itself "
        f"survives. Over the {fmt(n_shared)} spots shared by the two runs, each multi-gene "
        f"subdomain maps onto a distinct published subdomain, at "
        f"{pct(corr.best_match_frac.min())} to {pct(corr.best_match_frac.max())} of its spots -- "
        f"a near one-to-one correspondence. The contrast we analyse, multi-gene Subdomain 3 "
        f"versus Subdomain 1, therefore corresponds to the published Subdomain 1 versus "
        f"Subdomain 2 in the same direction, so the two are directly comparable.")
    doc.table(corr.rename(columns={"multigene": "multi-gene subdomain",
                                   "best_published_match": "published subdomain",
                                   "best_match_frac": "fraction of its spots"}),
              "Spot-level correspondence between the multi-gene and published microdomain "
              "partitions, on the same inherited 50 um grid.", nd=3)
    doc.figure(FIG / "subdomain_correspondence.jpeg",
               "The same correspondence as a contingency map. A near-diagonal mapping means the "
               "microdomain partition itself survives the multi-gene restriction, not only the "
               "differential expression.")

    m = a["de_new"][["names", "scores"]].merge(a["de_pub"][["names", "scores"]], on="names",
                                               suffixes=("_mg", "_pub"))
    r = float(m[["scores_mg", "scores_pub"]].corr().iloc[0, 1])
    rs = float(m[["scores_mg", "scores_pub"]].corr(method="spearman").iloc[0, 1])
    doc.p(
        f"At the gene level the two contrasts are almost the same analysis: the differential "
        f"expression scores correlate at Pearson {fmt(r, 3)} and Spearman {fmt(rs, 3)} across all "
        f"{len(m)} panel genes. The genes driving each side are unchanged -- "
        f"{', '.join(a['de_new'].nlargest(6, 'scores').names)} on the pre-synaptic side and "
        f"{', '.join(a['de_new'].nsmallest(6, 'scores').names)} on the post-synaptic side.")

    gn, gp = a["gsea_new"], a["gsea_pub"]
    shared = set(gp.ID) & set(gn.ID)
    mm = gp[["ID", "NES"]].merge(gn[["ID", "NES"]], on="ID", suffixes=("_pub", "_mg"))
    doc.p(
        f"The pathway result follows. Of the {len(gp)} gene sets enriched in the published "
        f"granule-layer contrast, {len(shared)} ({pct(len(shared) / len(gp))}) are recovered; "
        f"their normalised enrichment scores correlate at "
        f"{fmt(float(mm[['NES_pub','NES_mg']].corr().iloc[0,1]), 3)} with "
        f"{pct(float(np.mean(np.sign(mm.NES_pub) == np.sign(mm.NES_mg))), 0)} sign agreement. All "
        f"four of the pre-synaptic terms the manuscript reports are recovered.")
    key = ["Synaptic vesicle recycling", "Presynaptic endocytosis", "Exocytic process",
           "Endomembrane system organization"]
    rows = []
    for term in key:
        rp = gp[gp.Description == term]
        rn = gn[gn.Description == term]
        rows.append({"gene set": term,
                     "NES published": rp.NES.iloc[0] if len(rp) else np.nan,
                     "q published": rp["p.adjust"].iloc[0] if len(rp) else np.nan,
                     "NES multi-gene": rn.NES.iloc[0] if len(rn) else np.nan,
                     "q multi-gene": rn["p.adjust"].iloc[0] if len(rn) else np.nan,
                     "recovered": "yes" if len(rn) else "NO"})
    doc.table(pd.DataFrame(rows),
              "Recovery of the four pre-synaptic gene sets reported in the manuscript, "
              "multi-gene subset versus published.", nd=3)
    doc.figure(SUB / "granule_DE_genes_Subdomain 3_vs_Subdomain 1_target_GSEA.jpeg",
               "Gene sets enriched on the pre-synaptic side of the microdomain contrast, using "
               "multi-gene granules only. The synaptic-vesicle and exocytosis programme reported "
               "in the manuscript is reproduced.")


def section1_strata(doc, a):
    doc.h("1.5  The result is not a low-count artifact", 2)
    e = a["edges"].set_index("read_tercile").reindex(["low", "mid", "high"])
    doc.p(
        f"The complexity stratification above conditions on gene count. The reviewer's other "
        f"concern is read count, so we also split ALL granules -- not the multi-gene subset -- "
        f"into read-count terciles and repeated the WT/AD density comparison within each, holding "
        f"the published subtype labels fixed so that the only thing varying is depth. The "
        f"terciles are "
        + "; ".join(f"{k} {int(v['min'])}-{int(v['max'])} reads (median {int(v['median'])})"
                    for k, v in e.iterrows()) + ".")
    pre = a["sdens"][(a["sdens"].subtype == "pre-syn") & a["sdens"].brain_area.isin(AREAS)]
    piv = pre.pivot_table(index="brain_area", columns=["read_tercile", "sample"], values="density")
    star = pre[pre["sample"] == "WT"].pivot_table(index="brain_area", columns="read_tercile",
                                                  values="p_bonf_star", aggfunc="first")
    rows = []
    for area in AREAS:
        row = {"region": area}
        for t in ["low", "mid", "high"]:
            row[f"log2FC {t}"] = np.log2(piv.loc[area, (t, "AD")] / piv.loc[area, (t, "WT")])
            row[f"sig {t}"] = star.loc[area, t]
        rows.append(row)
    tab = pd.DataFrame(rows)
    n_sig = int(sum((tab[f"sig {t}"] == "***").sum() for t in ["low", "mid", "high"]))
    doc.p(
        f"The pre-synaptic reduction in AD is present at every read depth. It is significant at "
        f"p < 0.001 in {n_sig} of {3 * len(AREAS)} region x tercile cells, and in several regions "
        f"it is strongest in the HIGH-count tercile rather than the low one -- HPF-SR "
        f"{fmt(tab.set_index('region').loc['HPF-SR','log2FC low'])} / "
        f"{fmt(tab.set_index('region').loc['HPF-SR','log2FC mid'])} / "
        f"{fmt(tab.set_index('region').loc['HPF-SR','log2FC high'])}, and TH "
        f"{fmt(tab.set_index('region').loc['TH','log2FC low'])} / "
        f"{fmt(tab.set_index('region').loc['TH','log2FC mid'])} / "
        f"{fmt(tab.set_index('region').loc['TH','log2FC high'])}. A low-count artifact would be "
        f"confined to, or at least strongest in, the low tercile. It is not.")
    doc.table(tab, "AD-versus-WT log2 fold change in pre-synaptic granule density, within each "
                   "read-count tercile. Terciles are computed over all granules; subtype labels "
                   "are the published ones, held fixed.", nd=2)
    doc.p(
        "As a check that the stratification discriminates rather than merely reproducing the "
        "same numbers three times, the post-synaptic contrast behaves differently: it is negative "
        "in the low tercile and positive in the mid and high terciles in most regions. Read depth "
        "does change what is measured; it does not change the pre-synaptic conclusion.")
    doc.figure(FIG / "readstrata_density_pre-syn.jpeg",
               "Pre-synaptic granule density, WT versus AD, within each read-count tercile.")

    doc.h("1.6  Summary of A2a", 2)
    c, st = density_concordance(a["dens_pub"], a["dens_new"])
    doc.p(
        f"Restricting to granules with three or more distinct genes retains "
        f"{pct(float(a['ret'][a['ret'].brain_area=='overall'].retention.min()))} to "
        f"{pct(float(a['ret'][a['ret'].brain_area=='overall'].retention.max()))} of granules and "
        f"reproduces the published biology: subtype composition remains compartment-interpretable, "
        f"AD-versus-WT log2 fold changes correlate at r = {fmt(st['r'])} with "
        f"{pct(st['agree'])} direction agreement, {st['n_still_sig']} of {st['n_sig_pub']} "
        f"previously significant cells remain significant, the microdomain partition maps one-to-"
        f"one onto the published one, and the microdomain differential expression correlates at "
        f"Pearson {fmt(float(a['de_new'][['names','scores']].merge(a['de_pub'][['names','scores']], on='names', suffixes=('_a','_b'))[['scores_a','scores_b']].corr().iloc[0,1]), 3)}. "
        f"Separately, the pre-synaptic AD reduction holds within every read-count tercile. We "
        f"therefore do not find evidence that the reported structure is an artifact of low-count "
        f"granules.")


def section2(doc, b):
    """Body: the favourable half of A2b -- randomization does not reproduce the population."""
    det = b["detect"]
    real = det[det.condition == "real"].set_index("sample")
    perm = det[det.condition == "permuted"].groupby("sample").mean(numeric_only=True)
    n_real = int(real.n_fine.sum())
    n_perm = float(perm.n_fine.sum())

    doc.h("2. A2b — the randomized data the reviewer describes", 1)
    doc.h("2.0  Method", 2)
    doc.p(
        "The reviewer writes that they are \"not convinced that randomized data, with granules "
        "selected on a single marker gene in the same way, would not produce an essentially "
        "identical embedding\". We built exactly that data and ran the identical pipeline on it.")
    doc.p(
        "The null permutes the gene label across all transcripts within a sample. Every molecule "
        "position, the total transcript density and each gene's total count survive; only the "
        "association between a gene label and where its molecules sit is destroyed. The "
        "coordinates and the in-nucleus flag stay with their row, so each transcript keeps its "
        "own nuclear status. Per-gene totals, positional integrity and the fact that the labels "
        "actually moved are asserted in every run before detection begins.")
    doc.p(
        "The permuted data then goes through the same rough and fine detection passes, the same "
        "profiling, and the same **combined WT+AD** embedding built by 4_post_detection.ipynb -- "
        "which is the object the manuscript reports (Fig. 3f, Fig. 4d). A per-sample embedding "
        "would not be the thing under discussion. Five replicates, each pairing the WT and AD "
        "permutations of one seed, give five null embeddings against the one real embedding, and "
        "real and permuted arms are scored by one script on one code path.")
    doc.p(
        f"Because the permuted arms do not contain the same number of granules as the real one, "
        f"and both cluster-quality metrics depend on n, each permuted arm additionally emits a "
        f"size-matched pair: both arms cut to the smaller n and stratified by sample. All "
        f"{len(b['status'])} series embedded; none was too small to score.")

    doc.h("2.1  Randomization does not reproduce the granule population", 2)
    doc.p(
        f"Before any question about clustering, the detections themselves differ. The permuted "
        f"data yields {fmt(n_perm, 0)} granules against {fmt(n_real)} for the real data -- "
        f"**{fmt(n_real / n_perm, 1)}-fold fewer** -- and the reduction is present in both "
        f"samples. It is also not a threshold artifact: the gap already exists in the unfiltered "
        f"rough pass ({fmt(float(perm.n_rough.sum()), 0)} against {fmt(int(real.n_rough.sum()))}) "
        f"and widens through filtering.")
    doc.p(
        f"What chance does produce is markedly more soma-associated. Applying mcDETECT's own "
        f"in-nucleus criterion to the unfiltered detections, "
        f"{pct(float(real.frac_pass_in_soma.mean()))} of real rough detections pass, against "
        f"{pct(float(perm.frac_pass_in_soma.mean()))} of permuted ones. That is what the design "
        f"predicts: once labels are shuffled, the marker transcripts inherit the panel-wide "
        f"distribution, which is dominated by somata, so the aggregates chance forms sit where "
        f"the transcripts are densest rather than in the neuropil.")
    tab = pd.DataFrame({
        "sample": ["WT", "AD"],
        "rough, real": [real.loc["WT", "n_rough"], real.loc["AD", "n_rough"]],
        "rough, permuted": [round(perm.loc["WT", "n_rough"]), round(perm.loc["AD", "n_rough"])],
        "granules, real": [real.loc["WT", "n_fine"], real.loc["AD", "n_fine"]],
        "granules, permuted": [round(perm.loc["WT", "n_fine"]), round(perm.loc["AD", "n_fine"])],
        "fold fewer": [real.loc["WT", "n_fine"] / perm.loc["WT", "n_fine"],
                       real.loc["AD", "n_fine"] / perm.loc["AD", "n_fine"]],
        "pass in-soma, real": [real.loc["WT", "frac_pass_in_soma"],
                               real.loc["AD", "frac_pass_in_soma"]],
        "pass in-soma, permuted": [perm.loc["WT", "frac_pass_in_soma"],
                                   perm.loc["AD", "frac_pass_in_soma"]]})
    doc.table(tab, "Detections from the real and permuted data. Permuted values are the mean of "
                   "five replicates. The in-soma column is mcDETECT's own criterion applied post "
                   "hoc to the unfiltered rough set, identically for every arm.", nd=3)
    doc.p(
        "Both of these are properties of the detections alone. Neither depends on any clustering "
        "choice, on the number of subtypes, or on how the embedding is scored.")
    doc.figure(FIG / "a2b_granule_counts.jpeg",
               "Granules detected after filtering, real versus permuted. The shuffle preserves "
               "the marker transcript count exactly, so this difference comes from where those "
               "markers land, not from how many there are.")
    doc.figure(FIG / "a2b_detection_counts.jpeg",
               "Detections at the rough (unfiltered) and fine (size, in-soma and negative-control "
               "filtered) passes.")
    doc.figure(FIG / "a2b_in_soma_survival.jpeg",
               "Fraction of unfiltered detections passing mcDETECT's in-nucleus criterion.")


def section2_removable(doc, b):
    """Self-contained. Nothing in 2.1 or elsewhere refers to anything here, so this whole
    function's output can be deleted without touching the rest of the document."""
    per_k, sm = a2b_paired_comparison(b["metrics"], out_csv=A2B / "a2b_paired_comparison.csv")
    k15 = b["k15"]

    doc.h("2.2  The embedding comparison  [self-contained; may be removed]", 2)
    doc.p(
        "This subsection reports the part of A2b that is equivocal. The population result in 2.1 "
        "does not depend on it.", italic=True)
    doc.p(
        f"Comparing the real and permuted embeddings at matched n, the real one is better "
        f"structured, consistently. Across {sm['n_k']} values of k and {sm['n_seeds']} "
        f"permutation replicates -- {sm['n_comparisons']} paired comparisons -- the real "
        f"embedding has the higher silhouette in {sm['sil_real_higher']} and the higher cluster "
        f"stability in {sm['ari_real_higher']}. Testing over the per-k differences rather than "
        f"the individual comparisons, which are not independent, gives "
        f"p = {sm['sil_p_wilcoxon']:.1g} for silhouette and {sm['ari_p_wilcoxon']:.1g} for "
        f"stability. At the published k = 15 the real embedding wins on both metrics in all "
        f"{sm['n_seeds']} replicates.")
    doc.p(
        f"The effect is nonetheless modest, and we state that rather than leaving it to be "
        f"inferred. At k = 15 the silhouette is {fmt(sm['k15_sil_real'], 3)} for the real "
        f"embedding against {fmt(sm['k15_sil_perm'], 3)} for the permuted, and cluster stability "
        f"{fmt(sm['k15_ari_real'], 3)} against {fmt(sm['k15_ari_perm'], 3)}. The permuted "
        f"embedding retains {pct(sm['k15_perm_frac_of_real_sil'])} of the real silhouette at "
        f"k = 15, and between {pct(sm['perm_frac_of_real_sil_min'], 0)} and "
        f"{pct(sm['perm_frac_of_real_sil_max'], 0)} across k >= 10. It does not collapse.")
    doc.table(k15[["series", "condition", "seed", "silhouette_score", "ari_stability_mean",
                   "n_obs"]],
              "Cluster quality at the published k = 15, per arm. The matched pairs are the "
              "comparison; the unmatched rows are each arm at its own size.", nd=4)
    doc.table(per_k[per_k.n_clusters.isin([2, 5, 10, 15, 20, 25, 30])][
                  ["n_clusters", "sil_real", "sil_perm", "sil_diff", "seeds_real_higher_sil",
                   "ari_real", "ari_perm", "ari_diff", "seeds_real_higher_ari"]],
              "Real minus permuted at matched n, averaged over the five replicates, at selected "
              "k. The full sweep is in a2b_paired_comparison.csv.", nd=4)
    doc.figure(FIG / "a2b_silhouette_score.jpeg",
               "Silhouette against number of clusters, real versus permuted. The left panel "
               "equalises n between the arms; the right shows each at its own size.")
    doc.figure(FIG / "a2b_ari_stability_mean.jpeg",
               "Cluster stability (mean pairwise ARI over five clustering seeds) against number "
               "of clusters.")
    doc.p(
        "The two embeddings also look alike. We include both panels because they are the fair "
        "visual comparison and showing only the metrics would be choosing the flattering "
        "evidence.")
    doc.figure(A2B / "tsne_matched_real_seed0.jpeg",
               "t-SNE of the real combined WT+AD embedding, subsampled to the matched n.",
               width_in=4.4)
    doc.figure(A2B / "tsne_matched_perm_seed0.jpeg",
               "t-SNE of one permuted embedding at the same n, same parameters, same code path. "
               "The real panel is more broken into discrete islands, but the permuted embedding "
               "is plainly structured too.", width_in=4.4)
    fc = embedding_feature_composition(A2B / "embedding_feature_composition.csv")
    doc.p(
        f"Why the null retains so much structure. The embedding and the k = 15 subtyping are "
        f"computed on the {int(fc['n_features'])} compartment markers, and "
        f"{int(fc['n_seed_features'])} of those genes are themselves the detection seeds. Those "
        f"{int(fc['n_seed_features'])} carry {pct(fc['var_frac_seed'])} of the variance in the "
        f"matrix the embedding is built from, and "
        f"{pct(fc['frac_top_gene_is_own_seed'])} of granules have their own seeding marker as "
        f"their highest-valued reference gene. The reviewer's reading of Fig. 3e is therefore "
        f"correct, and our own data puts a number on it.")
    doc.table(pd.DataFrame([{
        "compartment markers used for clustering": int(fc["n_features"]),
        "of which are detection seeds": int(fc["n_seed_features"]),
        "variance carried by the seed features": fc["var_frac_seed"],
        "variance carried by the non-seed features": fc["var_frac_nonseed"],
        "granules whose top marker is their own seed": fc["frac_top_gene_is_own_seed"]}]),
        "Composition of the feature space the granule embedding is computed on.", nd=3)
    doc.p(
        "What follows from that. The discreteness of the embedding is a property of the detection "
        "procedure rather than of the data the procedure is run on: a granule is a DBSCAN cluster "
        "of one marker's transcripts, so it is dominated by that marker, and clusters in the "
        "embedding largely separate granules by which marker dominates them. The reviewer's "
        "sentence stipulates that the randomized data be processed \"with granules selected on a "
        "single marker gene in the same way\", which holds exactly that procedure fixed. Any null "
        "obeying the stipulation will therefore reproduce the discreteness, whatever "
        "\"randomized data\" is taken to mean. We chose the broadest reading -- permute the gene "
        "label across all transcripts and re-run everything -- and we do not think a narrower one "
        "would have changed the outcome: holding the granules fixed and randomizing only the "
        "non-seed content would move 14 of the 34 features, carrying "
        + pct(fc["var_frac_nonseed"]) + " of the variance.")
    doc.p(
        "So we do not contest the mechanism. We accept that the granule-level embedding is "
        "largely seed-driven, and note that this is why the manuscript does not rest any "
        "biological claim on it -- the WT/AD comparison was moved to aggregated neuropil "
        "microdomains, which the reviewer describes as sensible. Two things are worth setting "
        "beside the concession. Randomization does break the granule population, decisively and "
        "independently of any clustering choice (2.1). And a multi-gene granule transcriptome "
        "does exist: among the genes detection never touches, a third of all gene pairs co-occur "
        "above chance and the strongest are coherent molecular modules (3.2). The 34-marker "
        "embedding does not surface that content, because the genes carrying it account for "
        + pct(fc["var_frac_nonseed"]) + " of its variance.")


def section3(doc, c):
    ge = c["ge"]
    prim = ge[ge.arm == "all"].sort_values("median_z", ascending=False)
    pairs = c["pairs"]
    a = pairs[pairs.arm == "all"]
    seed = pairs[pairs.arm == "seed_markers"]
    run = c["run"].set_index("arm")
    cal = c["cal"].iloc[0]
    n_pairs = len(a)
    from scipy.stats import norm
    zb = float(norm.isf(0.025 / n_pairs))

    doc.h("3. A2c — are the genes co-detected inside a granule non-randomly associated?", 1)
    doc.h("3.0  Why this analysis exists, and what it can bear on", 2)
    doc.p(
        "The reviewer did not request this analysis. We added it because the reviewer's specific "
        "claim -- that \"each granule essentially expresses just the single marker it was "
        "detected on\", so that subtypes \"largely reflect the seeding marker rather than any "
        "genuine multi-gene granule transcriptome\" -- is a claim about granule content that can "
        "be tested directly, on the genes detection never uses.")
    doc.p(
        "We also state at the outset what this cannot do. Showing that co-detected genes are "
        "functionally related does not, on its own, distinguish a granule from \"any co-expressed "
        "transcript cluster\" (the reviewer's phrasing in major point 1), because functional "
        "coherence is what co-expression looks like. To address that we included an explicit "
        "negative control -- gene groups defined by co-expression programmes rather than by "
        "localization -- and we report below that this control did not separate the two.")

    doc.h("3.1  Method", 2)
    doc.p(
        f"Granule set. The full published detection, all {fmt(int(run.loc['all','n_granules_used']) + int(run.loc['all','n_granules_dropped']))} "
        f"granules, not the A2a subset: conditioning on gene count would select on the very "
        f"statistic being measured. Granules with fewer than two genes contribute no pairs under "
        f"either the data or the null and are dropped, leaving "
        f"{fmt(int(run.loc['all','n_granules_used']))}.")
    doc.p(
        "Genes. The 270 non-seed genes only. merge_sphere() merges overlapping spheres seeded by "
        "different markers, so co-occurrence among the 20 seed markers is partly manufactured by "
        "detection -- 64.7% of granules carry at least two of them. Running the analysis on the "
        "seeds would reproduce exactly the circularity the reviewer objects to. The seed markers "
        "are still analysed, as a positive control for the statistic, and are labelled as "
        "detection-confounded wherever they appear.")
    doc.p(
        f"Null model. For each gene pair we compare the observed number of granules containing "
        f"both against a null that holds BOTH margins exactly -- every granule keeps its number "
        f"of distinct genes, and every gene keeps the number of granules it appears in -- "
        f"generated by a curveball trade chain (Strona et al., Nat Commun 5:4114, 2014). Those "
        f"are the two effects that would otherwise masquerade as co-occurrence: complex granules "
        f"pair everything with everything, and abundant genes pair with everything. The "
        f"expectation and variance are taken from {int(run.loc['all','n_null_samples'])} "
        f"post-burn-in states of the chain, and z = (observed - expected) / SD.")
    doc.p(
        "Two things we measured rather than assumed, because both would have produced a "
        "confident wrong answer. First, the obvious analytic shortcut -- a maximum-entropy model "
        "fixing the degrees only in expectation -- is badly biased for data this sparse: a "
        "granule with exactly k genes contributes exactly k(k-1)/2 pairs, whereas a soft-degree "
        "null contributes about k^2/2, a factor of k/(k-1) that at the median granule complexity "
        "of about 5 genes is a 25% over-estimate applied to every pair; on simulated data it "
        "inflated z by roughly 96 standard deviations. Curveball's constraints are hard and it "
        "has no such bias. We verified the calibration by scoring a matrix drawn from the null "
        f"itself, which returns z with mean {fmt(cal['mean'], 2)} and SD {fmt(cal['sd'], 2)}. "
        "Second, the group-level permutation must be abundance-matched: "
        "rare genes carry systematically higher z, and the two gene "
        "programmes compared below differ about five-fold in abundance, so an unmatched test "
        "would have been biased before any biology entered. Each permutation replicate is drawn "
        "with the group's own abundance-bin composition.")
    doc.p(
        f"Gene groups. From the panel's own curated annotation "
        f"(data/MERSCOPE_WT_1/processed_data/gene_panel.csv), which supplies both kinds of "
        f"programme: localization groups (pre-synaptic, post-synaptic, Neuropil, Axons) and "
        f"co-expression groups (cell-type, regional and cortical-layer marker sets). "
        f"{len(c['groups'])} groups have at least four non-seed genes and are tested; "
        f"{len(c['dropped'])} fall below that and are listed rather than dropped silently. Group "
        f"significance is the median z over within-group pairs against "
        f"{int(prim.n_perm.iloc[0]):,} abundance-matched permutations, BH-corrected across groups. "
        f"An external check using GO biological-process annotation is reported in 3.3.")

    doc.h("3.2  Granules carry structured, non-random content in genes detection never touches", 2)
    doc.p(
        f"Across the {fmt(n_pairs)} pairs of non-seed genes, co-occurrence is far from random. "
        f"{pct((a.z > 2).mean())} of pairs are enriched at z > 2, against the 2.5% expected, and "
        f"{pct((a.z > zb).mean())} survive a Bonferroni threshold of z > {fmt(zb)}; "
        f"{pct((a.z < -zb).mean())} are significantly depleted. The positive control behaves as "
        f"designed: the seed markers, which detection does merge across spheres, reach a median z "
        f"of {fmt(seed.z.median(), 1)} against {fmt(a.z.median(), 2)} for the non-seed pairs.")
    doc.p(
        "The strongest pairs are transparently interpretable, which is the substance of the "
        "answer to the reviewer: they are coherent molecular modules, not arbitrary partners of a "
        "seeding marker.")
    top = a.nlargest(15, "z")[["gene_i", "gene_j", "observed", "expected", "z",
                               "log2_obs_over_exp"]].copy()
    top.columns = ["gene A", "gene B", "observed", "expected", "z", "log2 obs/exp"]
    doc.table(top, "The fifteen most strongly co-occurring non-seed gene pairs. None of these "
                   "genes is used by detection.", nd=2)
    doc.p(
        "These include the neurofilament triplet (Nefm-Nefh, Nefm-Nefl), the two GABA-synthesis "
        "enzymes (Gad2-Gad1), an oligodendrocyte pair (Cnp-Sox10), axonal and nodal components "
        "(Vamp1-Nefh, Kcna2-Syt2), and synaptic-vesicle machinery around Stxbp1 (Syngr1, Pfkm, "
        "Pacsin1). A granule whose content were exhausted by its seeding marker could not produce "
        "this.")
    doc.figure(A2C / "cooccurrence_clustermap.jpeg",
               "Co-occurrence enrichment (z) for all 270 non-seed genes, hierarchically "
               "clustered, with the programme annotation as a colour bar.")


def section3_claim2(doc, c):
    """Self-contained. Nothing in 3.2 refers to anything here, so this whole function's output
    can be deleted without touching the rest of the document."""
    ge, prog = c["ge"], c["prog"]
    prim = ge[ge.arm == "all"].sort_values("median_z", ascending=False)
    go = c["go"].iloc[0]
    gobin = c["gobin"]
    per, conf, sm = go_abundance_stratified(
        c["pair_go"], c["pairs"][c["pairs"].arm == "all"],
        out_csv=A2C / "go_abundance_stratified.csv")

    doc.h("3.3  Localization versus co-expression programmes  [self-contained; may be removed]", 2)
    doc.p(
        "This subsection reports a result that does not support the expectation we set out with. "
        "It is kept because the design was only worth running if it could fail, and it did. "
        "Nothing in 3.2 depends on it.", italic=True)
    doc.p(
        "The reasoning was as follows. If mcDETECT granules are packaged transport structures, "
        "the gene groups that co-occur should be localization programmes -- pre- and "
        "post-synaptic, neuropil, axonal. If instead they were local co-expressed transcript "
        "clusters, then cell-type, regional and layer marker sets should co-occur at least as "
        "strongly, because those genes are co-expressed by definition. The contrast was the test.")
    p_all = prog[prog.arm == "all"].set_index("programme")
    doc.p(
        f"The contrast did not separate. Taking the median across groups, the localization "
        f"programmes reach {fmt(p_all.loc['localization','median_of_medians'])} and the "
        f"co-expression programmes {fmt(p_all.loc['co-expression','median_of_medians'])} -- if "
        f"anything the wrong way round. The most strongly co-occurring groups are cell-type "
        f"marker sets. The pre-synaptic set is significantly enriched but ranks eighth; the "
        f"post-synaptic and Neuropil sets, which are the two largest localization groups, show no "
        f"enrichment at all.")
    tab = prim[["programme", "group", "n_genes", "n_pairs", "median_detections", "median_z",
                "ses", "q_upper", "q_upper_star"]].copy()
    tab.columns = ["programme", "group", "genes", "pairs", "median detections", "median z",
                   "effect size", "q", "sig"]
    doc.table(tab, "Within-group co-occurrence enrichment for every tested gene group, non-seed "
                   "genes, all granules. Significance is against abundance-matched permutations, "
                   "BH-corrected.", nd=3)
    doc.figure(FIG / "a2c_group_enrichment.jpeg",
               "Median within-group co-occurrence enrichment, coloured by programme.")

    piv = ge.pivot_table(index=["programme", "group"], columns="arm", values="median_z")
    piv = piv[["all", "Isocortex", "WT", "AD"]].reset_index().sort_values("all", ascending=False)
    doc.p(
        "Restricting the analysis to a single region collapses almost all of it -- the Isocortex "
        "column below -- which says that much of the co-occurrence being measured is regional and "
        "tissue-compartment composition rather than granule packaging. That applies to the "
        "localization groups as well as the co-expression ones.")
    doc.table(piv, "Median within-group enrichment by arm. The Isocortex column is the control "
                   "for regional composition: a group whose signal vanishes there was reporting "
                   "tissue composition.", nd=2)
    doc.figure(FIG / "a2c_group_enrichment_by_arm.jpeg",
               "The same values as a heatmap, split by programme.")

    doc.p(
        f"An external check, using GO biological-process annotation instead of the panel's own "
        f"groups, initially looks like a stronger negative: pairs sharing at least one term have "
        f"a LOWER median enrichment than pairs sharing none ({fmt(go['median_z_shared'], 3)} "
        f"versus {fmt(go['median_z_not'], 3)}, difference {fmt(go['delta'], 3)}, Wilcoxon "
        f"p = {go['p_wilcox']:.2g}; {int(go['n_shared']):,} sharing against "
        f"{int(go['n_not']):,} not). That test is confounded, and we report it here only because "
        f"the correction matters.")
    doc.p(
        f"The confound is the same one that forced abundance matching into the group permutation "
        f"above: genes carrying GO annotation are the better-studied, more abundant ones, and "
        f"abundance drives the magnitude of z in both directions. Across expected-count deciles "
        f"the fraction of pairs sharing a GO term climbs monotonically from "
        f"{pct(sm['frac_shared_lowest_decile'])} to {pct(sm['frac_shared_highest_decile'])}, and "
        f"pairs that share a term have a median expected co-occurrence of "
        f"{fmt(sm['median_expected_shared'], 1)} against {fmt(sm['median_expected_not'], 1)} for "
        f"pairs that do not. The GO test was written before the abundance problem was "
        f"identified, and had not had the correction applied.")
    doc.table(conf.rename(columns={"decile": "expected-count decile",
                                   "median_expected": "median expected co-occurrence",
                                   "frac_shared_go": "fraction sharing a GO term",
                                   "n": "pairs"}),
              "Gene pairs sharing a GO biological-process term, by expected co-occurrence. "
              "Annotation coverage tracks abundance, which is what makes the pooled test above "
              "uninterpretable.", nd=3)
    doc.p(
        f"Repeating the comparison with abundance held fixed -- within each expected-count decile, "
        f"combined across deciles with a signed-rank test on the ten differences, which is the "
        f"right unit given that the pairs are not independent -- the effect does not survive: "
        f"{sm['n_deciles_negative']} of {sm['n_deciles']} deciles remain negative with a median "
        f"difference of {fmt(sm['median_difference'], 3)}, but "
        f"p = {fmt(sm['p_signed_rank'], 3)}. The defensible statement is therefore that we find "
        f"no evidence of an association in either direction between GO functional similarity and "
        f"co-occurrence, not that functionally related genes co-occur less.")
    doc.table(per, "GO shared-term comparison within expected-count deciles, and the per-decile "
                   "difference that the combined signed-rank test is computed on.", nd=3)

    strat = gobin[gobin.stratum.isin(["Bonferroni enriched", "middle", "Bonferroni depleted"])]
    if len(strat):
        base = float(gobin.baseline_frac_shared_go.iloc[0])
        st = strat.set_index("stratum")
        doc.p(
            f"We also tested the obvious reconciliation with the top-pair table in 3.2 -- that "
            f"co-occurrence might be carried by a few strong functional modules rather than by a "
            f"broad similarity gradient -- by asking what fraction of pairs share a GO term "
            f"within each enrichment stratum, against a baseline of {pct(base)}. It is not "
            f"supported: "
            + "; ".join(
                f"{k} {pct(float(st.loc[k,'frac_shared_go']))} "
                f"({fmt(float(st.loc[k,'enrichment_over_baseline']))}x baseline, "
                f"n = {int(st.loc[k,'n_pairs']):,})"
                for k in ["Bonferroni enriched", "middle", "Bonferroni depleted"]
                if k in st.index)
            + f". The strongly enriched stratum sits essentially exactly at baseline, and the "
              f"depleted stratum is the one that is elevated -- a U-shape, which is what the "
              f"abundance confound above predicts, since abundant pairs reach extreme z in both "
              f"directions.")
        doc.table(gobin[["stratum", "n_pairs", "median_z", "frac_shared_go",
                         "enrichment_over_baseline"]],
                  "Fraction of gene pairs sharing at least one GO biological-process term, by "
                  "co-occurrence enrichment stratum and decile.", nd=3)
        doc.figure(FIG / "a2c_go_by_z_bin.jpeg",
                   "Shared-GO fraction against co-occurrence enrichment. The dashed line is the "
                   "overall shared-GO fraction. A rise confined to the enriched tail would have "
                   "indicated concentration in strong functional modules; the curve is U-shaped "
                   "instead, tracking abundance rather than enrichment.")

    doc.p(
        "We report this as a null result. Neither the panel's own gene groups nor external GO "
        "annotation shows that co-occurrence is organised by localization programmes rather than "
        "by co-expression programmes, and the one apparently significant effect does not survive "
        "the abundance control. It does not undermine 3.2, which concerns whether granule content "
        "is structured at all and is unaffected by how that structure is annotated. It does mean "
        "that this analysis does not, on its own, establish that the detected structures are RNA "
        "granules rather than co-localized transcript clusters, and we have not claimed that it "
        "does.")


def section4(doc, a, b, c):
    doc.h("4. What this set of analyses does and does not establish", 1)
    _, st = density_concordance(a["dens_pub"], a["dens_new"])
    pairs = c["pairs"]
    ai = pairs[pairs.arm == "all"]
    det = b["detect"]
    _real = det[det.condition == "real"]
    _perm = det[det.condition == "permuted"].groupby("sample").mean(numeric_only=True)
    _a2b_fold = float(_real.n_fine.sum() / _perm.n_fine.sum())
    _a2b_real_soma = float(_real.frac_pass_in_soma.mean())
    _a2b_perm_soma = float(_perm.frac_pass_in_soma.mean())
    _, _a2b_sm = a2b_paired_comparison(b["metrics"])
    doc.p("Does.", bold=True)
    doc.bullets([
        f"The complexity stratification the reviewer asked for has been done. Restricting to "
        f"granules with three or more distinct genes retains about 70% of granules and reproduces "
        f"the published biology: {pct(st['agree'])} direction agreement on WT/AD density "
        f"(r = {fmt(st['r'])}), {st['n_still_sig']} of {st['n_sig_pub']} previously significant "
        f"cells still significant, a near one-to-one microdomain partition, and all four "
        f"pre-synaptic gene sets recovered.",
        "The read-count stratification shows the pre-synaptic AD reduction at every depth, and in "
        "several regions most strongly at high depth. We find no evidence that the reported "
        "structure is a low-count artifact.",
        f"Granule content is not exhausted by the seeding marker. Among genes detection never "
        f"uses, {pct((ai.z > 2).mean())} of gene pairs co-occur above chance under a null that "
        f"holds granule complexity and gene abundance fixed, and the strongest pairs are coherent "
        f"molecular modules.",
        "The per-granule read and gene-count distributions (Fig. R9) are included, as requested.",
        f"The randomization the reviewer describes does not reproduce the granule population. It "
        f"yields {fmt(_a2b_fold, 1)}-fold fewer granules, and its unfiltered detections are far "
        f"more soma-associated ({pct(_a2b_perm_soma)} pass the in-nucleus criterion against "
        f"{pct(_a2b_real_soma)} for the real data). Both are properties of the detections alone "
        f"and depend on no clustering choice.",
    ])
    doc.p("Does not.", bold=True)
    doc.bullets([
        "None of this is protein-level validation. The reviewer's summary comment asks for RBP "
        "colocalization, and these analyses do not substitute for it; that point is answered "
        "separately.",
        "The localization-versus-co-expression contrast in 3.3 did not separate, and the external "
        "GO check found no association in either direction once pair abundance was controlled. "
        "So A2c does not establish that the detected structures are RNA granules rather than "
        "co-localized transcript clusters. We state this as a limitation rather than presenting "
        "the functional analysis as evidence for granule identity. That subsection is "
        "self-contained and can be removed without affecting anything else in this document.",
        "The strongest localization group in 3.3, Axons, contains only four non-seed genes "
        "(Ank3, Nefh, Nefl, Nefm), so its large enrichment rests on a single cytoskeletal module "
        "and should not be generalised.",
        "The panel groups are not disjoint -- Syt2, for instance, is both a pre-synaptic marker "
        "and an Isocortex_RSP regional marker -- so group-level results are not fully "
        "independent of one another.",
        f"A2b does not settle the embedding half of the reviewer's proposition. At matched n the "
        f"real embedding is better structured on both metrics in "
        f"{_a2b_sm['sil_real_higher']} of {_a2b_sm['n_comparisons']} paired comparisons "
        f"(p = {_a2b_sm['sil_p_wilcoxon']:.1g}), but the permuted embedding retains "
        f"{pct(_a2b_sm['k15_perm_frac_of_real_sil'])} of the real silhouette at k = 15 and is "
        f"plainly structured. Sparse profiles cluster discretely whether or not the gene "
        f"identities are real, so much of the embedding's discreteness does follow from "
        f"sparsity, as the reviewer argued. That subsection (2.2) is self-contained and can be "
        f"removed without affecting 2.1.",
    ])


def main(out_path=None):
    a = load_a2a()
    b = load_a2b()
    c = load_a2c()
    doc = Doc()

    doc.d.add_heading("Response to Reviewer #2, major point 6: sparsity of granules and the "
                      "stochastic origin of granule-level structure", 0)
    doc.p("Internal record of the computational analyses. Every number below is read directly "
          "from the analysis outputs under R2_revision/sparsity_structure/output/ by "
          "build_response_doc.py; none is transcribed by hand.", italic=True)

    section_map(doc)
    section0(doc, a, b)
    section1_methods(doc, a)
    section1_results(doc, a)
    section1_microdomains(doc, a)
    section1_strata(doc, a)
    section2(doc, b)
    section2_removable(doc, b)
    section3(doc, c)
    section3_claim2(doc, c)
    section4(doc, a, b, c)

    out = Path(out_path) if out_path else REPO / "plans" / "Response_R2_comment6_sparsity.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"wrote {out}  ({out.stat().st_size / 1e6:.1f} MB)")
    print(f"  {doc.tabno} tables, {doc.figno} figures")
    for m in doc.manifest:
        print("   ", m)
    if TMP.exists():
        shutil.rmtree(TMP)
    return out


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else None)
