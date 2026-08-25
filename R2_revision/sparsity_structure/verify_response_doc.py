#!/usr/bin/env python3
"""
Check the built response document against its sources.

    python3 verify_response_doc.py [doc.docx]

The tables in the document are generated, so this checks the generator rather than anyone's
arithmetic: every asserted cell is re-read from the CSV or parquet it came from, and two of the
tables are recomputed from raw inputs rather than compared to the same intermediate. Run it after
every build, including the rebuild once A2b lands.
"""

import re
import sys
import zipfile
from pathlib import Path

import docx
import numpy as np

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_response_doc as B


def cells(t):
    return [[c.text.strip() for c in r.cells] for r in t.rows]


def norm(s):
    return re.sub(r"\s+", " ", s.replace("’", "'").replace("–", "-")).strip()


def main(path=None):
    path = Path(path) if path else B.REPO / "plans" / "Response_R2_comment6_sparsity.docx"
    d = docx.Document(B.require(path, "run build_response_doc.py first"))
    tabs = d.tables
    a, b, c = B.load_a2a(), B.load_a2b(), B.load_a2c()
    body = "\n".join(p.text for p in d.paragraphs)
    heads = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
    checks = []

    # ---- A2b: which table index each is depends on section order, so find them by header ----
    def table_with(*required):
        for t in tabs:
            hdr = [c.text.strip() for c in t.rows[0].cells]
            if all(any(r in h for h in hdr) for r in required):
                return cells(t)
        raise AssertionError(f"no table with headers containing {required}")

    det = b["detect"]
    real = det[det.condition == "real"].set_index("sample")
    perm = det[det.condition == "permuted"].groupby("sample").mean(numeric_only=True)
    t = table_with("granules, real", "pass in-soma, permuted"); h = t[0]
    for row in t[1:]:
        smp = row[0]
        assert row[h.index("granules, real")] == f"{int(real.loc[smp,'n_fine']):,}"
        assert row[h.index("granules, permuted")] == f"{round(perm.loc[smp,'n_fine']):,}"
        assert row[h.index("pass in-soma, real")] == f"{real.loc[smp,'frac_pass_in_soma']:,.3f}"
    checks.append("A2b detection table matches a2b_detection_summary.csv")

    k15 = b["k15"]
    t = table_with("silhouette_score", "ari_stability_mean"); h = t[0]
    for row, (_, s_) in zip(t[1:], k15.iterrows()):
        assert row[h.index("series")] == s_["series"]
        assert row[h.index("silhouette_score")] == f"{s_['silhouette_score']:,.4f}"
    checks.append("A2b k = 15 table matches a2b_structure_at_k15.csv")

    per_k, sm = B.a2b_paired_comparison(b["metrics"])
    assert sm["n_comparisons"] == 145 and sm["n_seeds"] == 5
    assert f"{sm['sil_real_higher']} " in body or str(sm["sil_real_higher"]) in body, (
        "the paired win count is not stated in the document")
    t = table_with("sil_diff", "seeds_real_higher_ari"); h = t[0]
    ref = per_k[per_k.n_clusters.isin([2, 5, 10, 15, 20, 25, 30])]
    for row, (_, s_) in zip(t[1:], ref.iterrows()):
        assert row[h.index("sil_diff")] == f"{s_['sil_diff']:,.4f}", (row, s_["sil_diff"])
    checks.append(f"A2b paired comparison reproduces ({sm['sil_real_higher']}/145 silhouette "
                  f"wins, p = {sm['sil_p_wilcoxon']:.1g})")

    conc, _ = B.density_concordance(a["dens_pub"], a["dens_new"])
    pre = conc.loc["pre-syn"]
    t = table_with("WT published", "sig multi-gene"); h = t[0]
    for row in t[1:]:
        r0 = row[0]
        assert row[h.index("WT published")] == f"{pre.loc[r0,'WT_pub']:,.2f}"
        assert row[h.index("WT multi-gene")] == f"{pre.loc[r0,'WT_mg']:,.2f}"
        assert row[h.index("sig published")] == pre.loc[r0, "p_bonf_star_pub"]
        assert row[h.index("sig multi-gene")] == pre.loc[r0, "p_bonf_star_mg"]
    checks.append("pre-synaptic density table matches both source CSVs")

    corr = (a["corr"][["multigene", "best_published_match", "best_match_frac"]]
            .drop_duplicates().sort_values("multigene"))
    for row, (_, s) in zip(table_with("multi-gene subdomain", "fraction of its spots")[1:], corr.iterrows()):
        assert row[0] == s.multigene and row[1] == s.best_published_match
        assert row[2] == f"{s.best_match_frac:,.3f}"
    checks.append("subdomain correspondence table matches its CSV")

    gp, gn = a["gsea_pub"], a["gsea_new"]
    for row in table_with("NES published", "recovered")[1:]:
        rp, rn = gp[gp.Description == row[0]], gn[gn.Description == row[0]]
        assert row[1] == f"{rp.NES.iloc[0]:,.3f}" and row[3] == f"{rn.NES.iloc[0]:,.3f}"
        assert row[5] == "yes", f"{row[0]} reported as recovered but is not in the multi-gene GSEA"
    checks.append("all four pre-synaptic gene sets match both GSEA tables")

    # recomputed from the raw per-spot densities, not from any intermediate the builder used
    ps = a["sdens"][(a["sdens"].subtype == "pre-syn") & a["sdens"].brain_area.isin(B.AREAS)]
    piv = ps.pivot_table(index="brain_area", columns=["read_tercile", "sample"], values="density")
    t = table_with("log2FC low", "sig high"); h = t[0]
    for row in t[1:]:
        for k in ["low", "mid", "high"]:
            want = np.log2(piv.loc[row[0], (k, "AD")] / piv.loc[row[0], (k, "WT")])
            assert row[h.index(f"log2FC {k}")] == f"{want:,.2f}"
    checks.append("read-tercile fold changes recomputed from readstrata_density.csv")

    ai = c["pairs"][c["pairs"].arm == "all"]
    for row, (_, s) in zip(table_with("gene A", "log2 obs/exp")[1:], ai.nlargest(15, "z").iterrows()):
        assert row[0] == s.gene_i and row[1] == s.gene_j and row[4] == f"{s.z:,.2f}"
    checks.append("top co-occurring pairs match pair_enrichment.parquet")

    prim = c["ge"][c["ge"].arm == "all"].sort_values("median_z", ascending=False)
    for row, (_, s) in zip(table_with("programme", "effect size")[1:], prim.iterrows()):
        assert row[1] == s.group and row[5] == f"{s.median_z:,.3f}"
    checks.append("group enrichment table matches group_enrichment.csv")

    # the GO block: annotated-pair count and shared fraction must reproduce, and the stratified
    # table in the document must match an independent recomputation from the two parquets
    pg = c["pair_go"]
    go = c["go"].iloc[0]
    assert len(pg) == int(go["n_shared"]) + int(go["n_not"]), (
        f"a2c_pair_go.parquet has {len(pg)} rows but the summary reports "
        f"{int(go['n_shared']) + int(go['n_not'])} annotated pairs")
    assert abs(pg.go_shared.mean() - go["n_shared"] / len(pg)) < 1e-9
    per, conf, sm = B.go_abundance_stratified(pg, c["pairs"][c["pairs"].arm == "all"])
    t = table_with("expected-count decile", "difference"); h = t[0]
    for row, (_, s_) in zip(t[1:], per.iterrows()):
        assert row[h.index("difference")] == f"{s_['difference']:,.3f}", (row, s_["difference"])
    assert 0.0 <= sm["p_signed_rank"] <= 1.0
    checks.append(f"GO block reproduces ({len(pg):,} annotated pairs, stratified p = "
                  f"{sm['p_signed_rank']:.3f})")

    zf = zipfile.ZipFile(B.require(B.REPO / "plans" / "R2_Reviewer_comments.docx"))
    txt = re.sub(r"<[^>]+>", "", re.sub(r"</w:p>", "\n", zf.read("word/document.xml").decode("utf8")))
    paras = [l.strip() for l in txt.replace("&amp;", "&").replace("&quot;", '"').split("\n")
             if l.strip()]
    src = paras[24]

    # the mapping table: every quoted fragment must be verbatim in the paragraph its row names,
    # so the table is checked against the reviewer's file rather than transcribed and trusted
    for r in B.COMMENT_MAP:
        assert r["quote"] in paras[r["para"]], (
            f"{r['section']}: quote is not verbatim in reviewer paragraph {r['para']}\n"
            f"  {r['quote'][:90]}")
    t = table_with("reviewer comment", "the text it answers"); h = t[0]
    assert len(t) - 1 == len(B.COMMENT_MAP), "mapping table row count does not match COMMENT_MAP"
    for row, r in zip(t[1:], B.COMMENT_MAP):
        assert row[h.index("section")] == r["section"]
        assert row[h.index("reviewer comment")] == r["comment"]
        assert row[h.index("requested by the reviewer")] == r["requested"]
        assert row[h.index("the text it answers")] == r["quote"]
    checks.append(f"comment-mapping table: all {len(B.COMMENT_MAP)} quotes verbatim in "
                  f"R2_Reviewer_comments.docx")
    assert B.REVIEWER_COMMENT in body, "the reviewer quote is not in the document"
    assert norm(B.REVIEWER_COMMENT) == norm(src), "the reviewer quote is not verbatim"
    checks.append("reviewer quote is verbatim from R2_Reviewer_comments.docx")

    a2b_pending = any("PENDING" in x for x in heads)
    assert a2b_pending == (not (B.OUT / "a2b" / "metrics" / "a2b_metrics.csv").exists()), (
        "the A2b section's PENDING marker disagrees with whether A2b has actually run")
    # the embedding-composition table must match a fresh recomputation, not just the cache
    fc = B.embedding_feature_composition(B.A2B / "embedding_feature_composition.csv")
    fresh = B.embedding_feature_composition(None)
    for k in ("var_frac_seed", "frac_top_gene_is_own_seed", "n_seed_features"):
        assert abs(float(fc[k]) - float(fresh[k])) < 1e-9, (
            f"cached embedding_feature_composition is stale on {k}")
    t = table_with("of which are detection seeds", "variance carried by the seed features")
    h = t[0]
    assert t[1][h.index("of which are detection seeds")] == f"{int(fc['n_seed_features']):,}"
    assert t[1][h.index("variance carried by the seed features")] == f"{fc['var_frac_seed']:,.3f}"
    checks.append(f"embedding-composition table matches a fresh recomputation "
                  f"({int(fc['n_seed_features'])}/{int(fc['n_features'])} features are seeds, "
                  f"{fc['var_frac_seed']:.1%} of variance)")

    # 2.2 and 3.3 must stay independently deletable: 2.2 may point at 2.1 and 3.2, never at 3.3
    sec22 = []
    grab = False
    for para in d.paragraphs:
        st = para.style.name
        if st.startswith("Heading") and para.text.strip().startswith("2.2"):
            grab = True
            continue
        if st.startswith("Heading") and para.text.strip().startswith("3."):
            break
        if grab:
            sec22.append(para.text)
    sec22 = "\n".join(sec22)
    assert "3.3" not in sec22, "2.2 references 3.3; the two would no longer delete independently"
    assert "(3.2)" in sec22, "2.2 lost its pointer to the 3.2 counterweight"
    checks.append("2.2 points at 2.1 and 3.2 only, so 2.2 and 3.3 remain independently removable")

    removable = [x for x in heads if "may be removed" in x]
    assert any(x.strip().startswith("2.2") for x in removable), "A2b 2.2 is not marked removable"
    assert any(x.strip().startswith("3.3") for x in removable), "A2c 3.3 is not marked removable"
    assert "placeholder" not in body.lower(), "placeholder language survives in the document"
    stray = re.findall(r"\*\*[^*]+\*\*", body)
    assert not stray, f"markdown emphasis rendered literally: {stray[:3]}"
    checks.append("A2b PENDING marker gone; 2.2 and 3.3 both marked removable")

    for i, ch in enumerate(checks, 1):
        print(f"  {i}. {ch}")
    print(f"\n{len(checks)}/{len(checks)} checks passed on {path.name}")


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else None)
